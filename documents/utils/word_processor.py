import mammoth
from docx import Document
from docx.shared import Inches
import io
import base64
from datetime import datetime
from django.utils import timezone


class WordProcessor:
    """Processeur pour les fichiers Word (.docx et .doc)"""

    def process(self, file_path, document_instance):
        """Traite un fichier Word en maintenant le formatage"""
        try:
            # Utiliser mammoth pour extraire le HTML avec style
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                formatted_content = result.value
                conversion_messages = result.messages

            # Utiliser python-docx pour extraire les métadonnées et le texte brut
            doc = Document(file_path)

            # Extraire le texte brut
            content = self._extract_text_content(doc)

            # Extraire les métadonnées
            core_props = doc.core_properties

            # Extraire les images
            images = self._extract_images(doc)

            # Analyser la structure du document
            structure_info = self._analyze_document_structure(doc)

            # Générer le CSS personnalisé
            css_styles = self._generate_word_css(doc)

            # Améliorer le HTML avec le CSS
            enhanced_html = self._enhance_html_with_css(formatted_content, css_styles)

            return {
                'content': content,
                'formatted_content': enhanced_html,
                'author': core_props.author or '',
                'creation_date': self._convert_datetime(core_props.created),
                'modification_date': self._convert_datetime(core_props.modified),
                'images': images,
                'format_info': {
                    'page_width': None,  # Difficile à obtenir de python-docx
                    'page_height': None,
                    'fonts_used': structure_info.get('fonts', []),
                    'has_images': len(images) > 0,
                    'has_tables': structure_info.get('has_tables', False),
                    'has_headers': structure_info.get('has_headers', False),
                    'has_footers': structure_info.get('has_footers', False),
                    'generated_css': css_styles
                }
            }

        except Exception as e:
            raise Exception(f"Erreur lors du traitement du document Word: {str(e)}")

    def _extract_text_content(self, doc):
        """Extrait le texte brut du document"""
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs if p.text.strip()])
                    row_text.append(cell_text)
                if any(row_text):
                    text_content.append(' | '.join(row_text))

        return '\n'.join(text_content)

    def _extract_images(self, doc):
        """Extrait les images du document Word"""
        images = []

        # Les images dans python-docx sont dans les relations du document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_data = image_part.blob

                    # Convertir en base64
                    image_base64 = base64.b64encode(image_data).decode()

                    images.append({
                        'data': image_data,
                        'base64': image_base64,
                        'name': rel.target_ref.split('/')[-1],
                        'width': None,  # Difficile à obtenir sans plus d'analyse
                        'height': None
                    })
                except Exception as e:
                    continue

        return images

    def _analyze_document_structure(self, doc):
        """Analyse la structure du document"""
        fonts = set()
        has_tables = len(doc.tables) > 0
        has_headers = len([s for s in doc.sections if s.header.paragraphs]) > 0
        has_footers = len([s for s in doc.sections if s.footer.paragraphs]) > 0

        # Analyser les polices utilisées
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts.add(run.font.name)

        return {
            'fonts': list(fonts),
            'has_tables': has_tables,
            'has_headers': has_headers,
            'has_footers': has_footers
        }

    def _generate_word_css(self, doc):
        """Génère le CSS pour reproduire le style Word"""
        css = """
        .word-document {
            max-width: 8.5in;
            margin: 0 auto;
            padding: 1in;
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.15;
            background: white;
        }

        .word-paragraph {
            margin: 0 0 8pt 0;
            text-align: left;
        }

        .word-heading1 {
            font-size: 16pt;
            font-weight: bold;
            margin: 12pt 0 3pt 0;
        }

        .word-heading2 {
            font-size: 14pt;
            font-weight: bold;
            margin: 10pt 0 2pt 0;
        }

        .word-heading3 {
            font-size: 12pt;
            font-weight: bold;
            margin: 8pt 0 2pt 0;
        }

        .word-table {
            border-collapse: collapse;
            width: 100%;
            margin: 8pt 0;
        }

        .word-table td, .word-table th {
            border: 1pt solid #000;
            padding: 4pt;
            text-align: left;
        }

        .word-image {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 8pt 0;
        }
        """

        # Ajouter des styles spécifiques pour les polices trouvées
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name and run.font.size:
                        font_name = run.font.name
                        font_size = run.font.size.pt if run.font.size else 12
                        css += f"""
                        .font-{font_name.replace(' ', '-').lower()} {{
                            font-family: '{font_name}', serif;
                            font-size: {font_size}pt;
                        }}
                        """
        except:
            pass

        return css

    def _enhance_html_with_css(self, html_content, css_styles):
        """Améliore le HTML avec le CSS personnalisé"""
        enhanced_html = f"""
        <style>
        {css_styles}
        </style>
        <div class="word-document">
        {html_content}
        </div>
        """

        # Remplacer certains éléments HTML pour correspondre aux classes CSS
        enhanced_html = enhanced_html.replace('<p>', '<p class="word-paragraph">')
        enhanced_html = enhanced_html.replace('<h1>', '<h1 class="word-heading1">')
        enhanced_html = enhanced_html.replace('<h2>', '<h2 class="word-heading2">')
        enhanced_html = enhanced_html.replace('<h3>', '<h3 class="word-heading3">')
        enhanced_html = enhanced_html.replace('<table>', '<table class="word-table">')
        enhanced_html = enhanced_html.replace('<img', '<img class="word-image"')

        return enhanced_html

    def _convert_datetime(self, dt):
        """Convertit un datetime en datetime Django aware"""
        if dt is None:
            return None

        if dt.tzinfo is None:
            return timezone.make_aware(dt)

        return dt

    def process_legacy_doc(self, file_path, document_instance):
        """Traite les anciens fichiers .doc (nécessite conversion préalable)"""
        # Pour les fichiers .doc, il faudrait utiliser un convertisseur externe
        # comme LibreOffice ou antiword, ou demander à l'utilisateur de convertir
        # le fichier en .docx

        try:
            # Tentative avec mammoth (peut fonctionner sur certains .doc)
            with open(file_path, "rb") as doc_file:
                result = mammoth.convert_to_html(doc_file)
                formatted_content = result.value

            # Extraire le texte basique
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(formatted_content, 'html.parser')
            content = soup.get_text()

            return {
                'content': content,
                'formatted_content': f'<div class="word-document">{formatted_content}</div>',
                'author': '',
                'creation_date': None,
                'modification_date': None,
                'images': [],
                'format_info': {
                    'fonts_used': [],
                    'has_images': False,
                    'has_tables': bool(soup.find_all('table')),
                    'has_headers': False,
                    'has_footers': False,
                    'generated_css': self._generate_word_css(None)
                }
            }

        except Exception as e:
            raise Exception(f"Erreur lors du traitement du fichier .doc: {str(e)}. "
                            "Veuillez convertir le fichier en .docx pour un meilleur support.")